"""文档处理器测试"""
import pytest
import os
from io import BytesIO
from docx import Document
from app.services.document_processor import DocumentProcessor


class TestDocumentProcessor:
    """测试文档处理器"""

    def test_extract_text_from_txt(self, tmp_path):
        """测试从文本文件提取文本"""
        # 创建测试文件
        test_file = tmp_path / "test.txt"
        test_content = "这是测试文本内容"
        test_file.write_text(test_content, encoding='utf-8')

        # 提取文本
        result = DocumentProcessor.extract_text_from_txt(str(test_file))
        assert result == test_content

    def test_extract_text_from_docx(self, tmp_path):
        """测试从Word文档提取文本"""
        # 创建测试Word文档
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")

        test_file = tmp_path / "test.docx"
        doc.save(str(test_file))

        # 提取文本
        result = DocumentProcessor.extract_text_from_docx(str(test_file))
        assert "第一段内容" in result
        assert "第二段内容" in result

    def test_extract_text_invalid_file(self):
        """测试处理无效文件"""
        with pytest.raises(FileNotFoundError):
            DocumentProcessor.extract_text("non_existent_file.txt")

    def test_extract_text_unsupported_format(self, tmp_path):
        """测试不支持的文件格式"""
        test_file = tmp_path / "test.unsupported"
        test_file.write_text("content")

        with pytest.raises(ValueError):
            DocumentProcessor.extract_text(str(test_file))

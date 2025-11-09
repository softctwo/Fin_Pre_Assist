""" 文档导出服务"""

import os
from datetime import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from loguru import logger

from app.core.config import settings
from app.models import Proposal


class ExportService:
    """文档导出服务"""

    @staticmethod
    def export_proposal_to_word(proposal: Proposal) -> str:
        """导出方案为Word文档"""
        try:
            # 创建Word文档
            doc = docx.Document()

            # 设置文档样式
            ExportService._set_word_styles(doc)

            # 添加标题
            title = doc.add_heading(proposal.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加文档信息
            doc.add_paragraph(f"客户名称: {proposal.customer_name}")
            doc.add_paragraph(f"创建时间: {proposal.created_at.strftime('%Y年%m月%d日')}")
            if proposal.customer_industry:
                doc.add_paragraph(f"所属行业: {proposal.customer_industry}")

            doc.add_paragraph()  # 空行
            doc.add_page_break()  # 分页

            # 1. 客户需求
            doc.add_heading("一、客户需求", level=1)
            ExportService._add_paragraph_with_indent(doc, proposal.requirements)
            doc.add_page_break()

            # 2. 执行摘要
            if proposal.executive_summary:
                doc.add_heading("二、执行摘要", level=1)
                ExportService._add_markdown_content(doc, proposal.executive_summary)
                doc.add_page_break()

            # 3. 解决方案概述
            if proposal.solution_overview:
                doc.add_heading("三、解决方案概述", level=1)
                ExportService._add_markdown_content(doc, proposal.solution_overview)
                doc.add_page_break()

            # 4. 技术实现方案
            if proposal.technical_details:
                doc.add_heading("四、技术实现方案", level=1)
                ExportService._add_markdown_content(doc, proposal.technical_details)
                doc.add_page_break()

            # 5. 项目实施计划
            if proposal.implementation_plan:
                doc.add_heading("五、项目实施计划", level=1)
                ExportService._add_markdown_content(doc, proposal.implementation_plan)

            # 保存文档
            filename = f"proposal_{proposal.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            filepath = os.path.join(settings.EXPORT_DIR, filename)
            doc.save(filepath)

            logger.info(f"Word文档导出成功: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"导出Word文档失败: {e}")
            raise

    @staticmethod
    def export_proposal_to_pdf(proposal: Proposal) -> str:
        """导出方案为PDF文档"""
        try:
            # 设置PDF文件路径
            filename = f"proposal_{proposal.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
            filepath = os.path.join(settings.EXPORT_DIR, filename)

            # 创建PDF文档
            doc = SimpleDocTemplate(filepath, pagesize=A4)

            # 获取样式
            styles = getSampleStyleSheet()

            # 尝试注册中文字体
            try:
                # 尝试注册常见的中文字体
                font_paths = [
                    '/System/Library/Fonts/PingFang.ttc',
                    '/System/Library/Fonts/STSong.ttc',
                    '/System/Library/Fonts/Hiragino Sans GB.ttc',
                    '/System/Library/Fonts/Hiragino Sans GB W3.ttc'
                ]

                chinese_font = None
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            chinese_font = 'ChineseFont'
                            logger.info(f"成功注册中文字体: {font_path}")
                            break
                        except Exception as font_error:
                            logger.warning(f"字体注册失败 {font_path}: {font_error}")
                            continue
            except Exception as e:
                logger.warning(f"中文字体注册过程出错: {e}")

            # 创建自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor='black',
                fontName=chinese_font
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20,
                textColor='black',
                fontName=chinese_font
            )

            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=6,
                leading=14,
                alignment=TA_JUSTIFY,
                textColor='black',
                fontName=chinese_font
            )

            # 构建文档内容
            story = []

            # 标题
            story.append(Paragraph(proposal.title, title_style))
            story.append(Spacer(1, 12))

            # 基本信息
            story.append(Paragraph(f"<b>客户名称:</b> {proposal.customer_name}", content_style))
            story.append(Paragraph(f"<b>创建时间:</b> {proposal.created_at.strftime('%Y年%m月%d日')}", content_style))
            if proposal.customer_industry:
                story.append(Paragraph(f"<b>所属行业:</b> {proposal.customer_industry}", content_style))
            story.append(Spacer(1, 20))

            # 客户需求
            story.append(Paragraph("一、客户需求", heading_style))
            # 处理文本，保留换行
            requirements_text = proposal.requirements.replace('\n', '<br/>').replace('\r\n', '<br/>')
            story.append(Paragraph(requirements_text, content_style))
            story.append(Spacer(1, 12))

            # 执行摘要
            if proposal.executive_summary:
                story.append(PageBreak())
                story.append(Paragraph("二、执行摘要", heading_style))
                exec_summary_text = proposal.executive_summary.replace('\n', '<br/>').replace('\r\n', '<br/>')
                story.append(Paragraph(exec_summary_text, content_style))
                story.append(Spacer(1, 12))

            # 解决方案概述
            if proposal.solution_overview:
                story.append(PageBreak())
                story.append(Paragraph("三、解决方案概述", heading_style))
                solution_text = proposal.solution_overview.replace('\n', '<br/>').replace('\r\n', '<br/>')
                story.append(Paragraph(solution_text, content_style))
                story.append(Spacer(1, 12))

            # 技术实现方案
            if proposal.technical_details:
                story.append(PageBreak())
                story.append(Paragraph("四、技术实现方案", heading_style))
                tech_text = proposal.technical_details.replace('\n', '<br/>').replace('\r\n', '<br/>')
                story.append(Paragraph(tech_text, content_style))
                story.append(Spacer(1, 12))

            # 项目实施计划
            if proposal.implementation_plan:
                story.append(PageBreak())
                story.append(Paragraph("五、项目实施计划", heading_style))
                impl_text = proposal.implementation_plan.replace('\n', '<br/>').replace('\r\n', '<br/>')
                story.append(Paragraph(impl_text, content_style))

            # 生成PDF
            doc.build(story)

            logger.info(f"PDF文档导出成功: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"导出PDF失败: {e}")
            raise

    @staticmethod
    def export_pricing_to_excel(proposal: Proposal) -> str:
        """导出报价单为Excel"""
        try:
            # 创建工作簿
            wb = Workbook()
            ws = wb.active
            ws.title = "报价单"

            # 设置列宽
            ws.column_dimensions["A"].width = 5
            ws.column_dimensions["B"].width = 30
            ws.column_dimensions["C"].width = 15
            ws.column_dimensions["D"].width = 15
            ws.column_dimensions["E"].width = 20

            # 定义样式
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=12)
            title_font = Font(bold=True, size=16)
            border = Border(
                left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin")
            )

            # 标题
            ws.merge_cells("B2:E2")
            title_cell = ws["B2"]
            title_cell.value = f"{proposal.customer_name} - 项目报价单"
            title_cell.font = title_font
            title_cell.alignment = Alignment(horizontal="center", vertical="center")

            # 基本信息
            ws["B4"] = "项目名称:"
            ws["C4"] = proposal.title
            ws["B5"] = "报价日期:"
            ws["C5"] = datetime.now().strftime("%Y年%m月%d日")
            ws["B6"] = "有效期:"
            ws["C6"] = "30天"

            # 表头
            row = 8
            headers = ["序号", "项目名称", "数量", "单价(元)", "小计(元)"]
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=row, column=col)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            # 报价项目
            items = [
                ("1", "软件许可费", "1", "200,000", "200,000"),
                ("2", "实施服务费", "1", "100,000", "100,000"),
                ("3", "培训费用", "1", "30,000", "30,000"),
                ("4", "年度维护费", "1", "50,000", "50,000"),
            ]

            row = 9
            total = 0
            for item in items:
                for col, value in enumerate(item, start=1):
                    cell = ws.cell(row=row, column=col)
                    cell.value = value
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                    cell.border = border
                row += 1
                if col >= 4:
                    total += int(value.replace(",", ""))

            # 合计
            row += 1
            ws.merge_cells(f"B{row}:D{row}")
            total_label = ws[f"B{row}"]
            total_label.value = "合计"
            total_label.font = Font(bold=True, size=12)
            total_label.alignment = Alignment(horizontal="right", vertical="center")
            total_label.border = border

            total_cell = ws[f"E{row}"]
            total_cell.value = f"{total:,}"
            total_cell.font = Font(bold=True, size=12, color="FF0000")
            total_cell.alignment = Alignment(horizontal="center", vertical="center")
            total_cell.border = border

            # 备注
            row += 3
            ws[f"B{row}"] = "备注:"
            ws[f"B{row}"].font = Font(bold=True)
            row += 1
            ws[f"B{row}"] = "1. 以上报价含税价"
            row += 1
            ws[f"B{row}"] = "2. 付款方式: 签约后30%，验收后70%"
            row += 1
            ws[f"B{row}"] = "3. 实施周期: 3个月"

            # 保存文件
            filename = f"quotation_{proposal.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
            filepath = os.path.join(settings.EXPORT_DIR, filename)
            wb.save(filepath)

            logger.info(f"Excel报价单导出成功: {filepath}")
            return filepath

        except Exception as e:
            logger.error(f"导出Excel失败: {e}")
            raise

    @staticmethod
    def _set_word_styles(doc):
        """设置Word文档样式"""
        # 设置正文样式
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)

    @staticmethod
    def _add_paragraph_with_indent(doc, text: str, indent: float = 0.5):
        """添加带缩进的段落"""
        paragraph = doc.add_paragraph(text)
        paragraph.paragraph_format.first_line_indent = Inches(indent)
        paragraph.paragraph_format.line_spacing = 1.5
        return paragraph

    @staticmethod
    def _add_markdown_content(doc, markdown_text: str):
        """将Markdown内容添加到Word文档"""
        # 简单处理：按行分割并添加
        lines = markdown_text.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 处理标题
            if line.startswith("###"):
                doc.add_heading(line.replace("###", "").strip(), level=3)
            elif line.startswith("##"):
                doc.add_heading(line.replace("##", "").strip(), level=2)
            elif line.startswith("#"):
                doc.add_heading(line.replace("#", "").strip(), level=1)
            # 处理列表
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0:2].replace(".", "").isdigit():
                doc.add_paragraph(line.split(".", 1)[1].strip(), style="List Number")
            # 普通段落
            else:
                ExportService._add_paragraph_with_indent(doc, line, indent=0)


# 全局实例
export_service = ExportService()
